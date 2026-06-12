import streamlit as st
from agent import generate_document
from db import save_document, get_history, delete_document
from docx import Document
import io

st.set_page_config(
    page_title="AI Document Generator",
    page_icon="📄",
    layout="wide"
)

# ── Custom CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
    .app-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 2rem; border-radius: 12px;
        margin-bottom: 1.5rem; text-align: center;
    }
    .app-header h1 { color: white; margin: 0; font-size: 2rem; }
    .app-header p  { color: rgba(255,255,255,0.85); margin: 0.4rem 0 0; font-size: 1rem; }

    div[data-testid="column"] button {
        width: 100%; border-radius: 10px !important;
        border: 2px solid transparent !important;
        padding: 0.6rem !important; font-size: 0.85rem !important;
        transition: all 0.2s ease !important;
    }
    /* Sidebar History Buttons */
    section[data-testid="stSidebar"] .stButton button {
        background: #f1f3f9 !important;
        color: #1a1a2e !important;
        border: 1px solid #d6dcf5 !important;
        border-radius: 12px !important;
        padding: 0.75rem !important;
        transition: all 0.25s ease !important;
        text-align: left !important;
    }

    /* Hover Effect */
    section[data-testid="stSidebar"] .stButton button:hover {
        background: linear-gradient(135deg, #667eea, #764ba2) !important;
        color: white !important;
        border-color: #667eea !important;
        transform: translateY(-2px);
        box-shadow: 0 6px 18px rgba(102,126,234,0.35) !important;
    }
            
    .selected-type {
        background: linear-gradient(135deg, #667eea, #764ba2);
        color: white; padding: 0.4rem 1rem; border-radius: 20px;
        font-size: 0.85rem; font-weight: 600;
        display: inline-block; margin-bottom: 1rem;
    }
    .word-count { color: #888; font-size: 0.8rem; text-align: right; margin-top: -0.5rem; }
    .output-card {
        background: #f8f9ff; border: 1px solid #e0e4ff;
        border-radius: 12px; padding: 1.5rem; margin-top: 1rem;
    }
    .hist-card {
        background: #fff; border: 1px solid #e0e4ff;
        border-radius: 10px; padding: 10px 12px; margin-bottom: 8px;
        cursor: pointer;
    }
    .hist-card:hover { border-color: #667eea; background: #f8f9ff; }
    .hist-badge {
        display: inline-block; font-size: 0.7rem; font-weight: 700;
        padding: 2px 8px; border-radius: 10px; margin-bottom: 4px;
        background: #667eea; color: white;
    }
    .hist-req  { font-size: 0.78rem; color: #444; line-height: 1.4; }
    .hist-time { font-size: 0.7rem; color: #999; margin-top: 4px; }

    section[data-testid="stSidebar"] { background: #ffffff !important; }
    section[data-testid="stSidebar"] * { color: #1a1a2e !important; }
    section[data-testid="stSidebar"] hr { border-color: #ddd !important; }
</style>
""", unsafe_allow_html=True)

# ── Session state ─────────────────────────────────────────────────────────────
if "doc_type"        not in st.session_state: st.session_state.doc_type = "BRD"
if "loaded_content"  not in st.session_state: st.session_state.loaded_content = None
if "loaded_doc_type" not in st.session_state: st.session_state.loaded_doc_type = None

# ── Header ────────────────────────────────────────────────────────────────────
st.markdown("""
<div class="app-header">
    <h1>📄 AI Document Generator</h1>
    <p>Generate enterprise-grade documents using LangGraph + Groq AI</p>
</div>
""", unsafe_allow_html=True)

# ── Sidebar ───────────────────────────────────────────────────────────────────
with st.sidebar:
    # ── Info section ──────────────────────────────────────────────────────────
    st.markdown("""
<div style="padding:0.5rem 0 0;">
<p style="font-size:0.95rem;font-weight:700;color:#1a1a2e;margin-bottom:0.6rem;">📖 Document Types</p>
<table style="width:100%;border-collapse:collapse;font-size:0.82rem;">
  <thead><tr>
    <th style="background:#e8eaff;color:#1a1a2e;padding:5px 8px;text-align:left;border:1px solid #ccd;">Type</th>
    <th style="background:#e8eaff;color:#1a1a2e;padding:5px 8px;text-align:left;border:1px solid #ccd;">Full Name</th>
  </tr></thead>
  <tbody>
    <tr><td style="padding:4px 8px;border:1px solid #dde;color:#1a1a2e;font-weight:600;">BRD</td><td style="padding:4px 8px;border:1px solid #dde;color:#333;">Business Requirements</td></tr>
    <tr style="background:#f8f9ff;"><td style="padding:4px 8px;border:1px solid #dde;color:#1a1a2e;font-weight:600;">BPD</td><td style="padding:4px 8px;border:1px solid #dde;color:#333;">Business Process</td></tr>
    <tr><td style="padding:4px 8px;border:1px solid #dde;color:#1a1a2e;font-weight:600;">FD</td><td style="padding:4px 8px;border:1px solid #dde;color:#333;">Functional Design</td></tr>
    <tr style="background:#f8f9ff;"><td style="padding:4px 8px;border:1px solid #dde;color:#1a1a2e;font-weight:600;">TD</td><td style="padding:4px 8px;border:1px solid #dde;color:#333;">Technical Design</td></tr>
    <tr><td style="padding:4px 8px;border:1px solid #dde;color:#1a1a2e;font-weight:600;">Custom</td><td style="padding:4px 8px;border:1px solid #dde;color:#333;">Any Document</td></tr>
  </tbody>
</table>
<hr style="border:none;border-top:1px solid #ddd;margin:0.8rem 0;"/>
<p style="font-size:0.95rem;font-weight:700;color:#1a1a2e;margin-bottom:0.5rem;">🔧 Tech Stack</p>
<div style="display:flex;flex-direction:column;gap:5px;">
  <div style="background:#f0f2ff;border-radius:7px;padding:5px 9px;font-size:0.8rem;color:#1a1a2e;">🦜 LangChain + LangGraph</div>
  <div style="background:#f0f2ff;border-radius:7px;padding:5px 9px;font-size:0.8rem;color:#1a1a2e;">⚡ Groq (LLaMA 3.3 70B)</div>
  <div style="background:#f0f2ff;border-radius:7px;padding:5px 9px;font-size:0.8rem;color:#1a1a2e;">🎈 Streamlit</div>
  <div style="background:#f0f2ff;border-radius:7px;padding:5px 9px;font-size:0.8rem;color:#1a1a2e;">📝 python-docx</div>
</div>
<hr style="border:none;border-top:1px solid #ddd;margin:0.8rem 0;"/>
</div>
""", unsafe_allow_html=True)

    # ── History section ───────────────────────────────────────────────────────
    st.markdown(
        '<p style="font-size:0.95rem;font-weight:700;color:#1a1a2e;margin-bottom:0.5rem;">📂 Document History</p>',
        unsafe_allow_html=True
    )

    history = get_history(limit=10)

    if not history:
        st.markdown(
            '<p style="font-size:0.78rem;color:#888;font-style:italic;">No documents yet. Generate one!</p>',
            unsafe_allow_html=True
        )
    else:
        for doc in history:
            short_req = doc["requirements"][:60] + "..." if len(doc["requirements"]) > 60 else doc["requirements"]
            col_card, col_del = st.columns([5, 1])

            with col_card:
                if st.button(
                    f"**{doc['doc_type']}** · {short_req}\n\n_{doc['created_at']}_",
                    key=f"hist_{doc['id']}",
                    use_container_width=True
                ):
                    st.session_state.loaded_content  = doc["content"]
                    st.session_state.loaded_doc_type = doc["doc_type"]
                    st.rerun()

            with col_del:
                if st.button("🗑️", key=f"del_{doc['id']}", help="Delete this document"):
                    delete_document(doc["id"])
                    if st.session_state.loaded_content == doc["content"]:
                        st.session_state.loaded_content  = None
                        st.session_state.loaded_doc_type = None
                    st.rerun()

    st.markdown(
        '<p style="font-size:0.72rem;color:#888;margin-top:0.8rem;">Built by Kesavan P</p>',
        unsafe_allow_html=True
    )

# ── Doc Type Selector ─────────────────────────────────────────────────────────
st.markdown("#### Step 1 — Choose Document Type")

doc_types = {
    "BRD":    ("📋", "Business Requirements"),
    "BPD":    ("🔄", "Business Process"),
    "FD":     ("⚙️", "Functional Design"),
    "TD":     ("🏗️", "Technical Design"),
    "Custom": ("✨", "Custom Document"),
}

cols = st.columns(5)
for i, (dtype, (icon, label)) in enumerate(doc_types.items()):
    with cols[i]:
        if st.button(
            f"{icon}\n\n**{dtype}**\n\n{label}",
            key=f"btn_{dtype}",
            use_container_width=True,
            type="primary" if st.session_state.doc_type == dtype else "secondary"
        ):
            st.session_state.doc_type = dtype
            st.rerun()

selected  = st.session_state.doc_type
icon_sel  = doc_types[selected][0]
label_sel = doc_types[selected][1]
st.markdown(f'<span class="selected-type">{icon_sel} {selected} — {label_sel}</span>', unsafe_allow_html=True)
st.markdown("---")

# ── Requirements Input ────────────────────────────────────────────────────────
st.markdown("#### Step 2 — Describe Your Requirements")

requirements = st.text_area(
    label="Requirements",
    label_visibility="collapsed",
    height=180,
    placeholder=(
        "Example: Build an e-commerce platform with user login, "
        "product catalog, shopping cart, and Razorpay payment gateway. "
        "Target: 10,000 daily users. Tech stack: React + FastAPI + PostgreSQL."
    )
)

word_count = len(requirements.split()) if requirements.strip() else 0
char_count = len(requirements)
_, col_wc = st.columns([4, 1])
with col_wc:
    color = "#27ae60" if word_count >= 20 else "#e67e22" if word_count >= 5 else "#e74c3c"
    st.markdown(
        f'<p class="word-count" style="color:{color}">📝 {word_count} words · {char_count} chars</p>',
        unsafe_allow_html=True
    )
if 0 < word_count < 5:
    st.caption("⚠️ Add more detail for a richer document (aim for 20+ words)")

st.markdown("---")

# ── Generate Button ───────────────────────────────────────────────────────────
st.markdown("#### Step 3 — Generate")

generate_btn = st.button(
    f"🚀 Generate {selected} Document",
    type="primary",
    use_container_width=True,
    disabled=(requirements.strip() == "")
)
if requirements.strip() == "":
    st.caption("⬆️ Enter your requirements above to enable generation")

# ── Generation Logic ──────────────────────────────────────────────────────────
if generate_btn:
    content = st.write_stream(generate_document(selected, requirements))
    save_document(selected, requirements, content)
    st.session_state.loaded_content  = content
    st.session_state.loaded_doc_type = selected
    st.rerun()

# ── Output area (shows both newly generated AND loaded from history) ──────────
if st.session_state.loaded_content:
    content   = st.session_state.loaded_content
    disp_type = st.session_state.loaded_doc_type or selected
    label_disp = doc_types.get(disp_type, ("📄", disp_type))[1]

    st.markdown("---")
    col_title, col_badge = st.columns([3, 1])
    with col_title:
        st.markdown(f"### 📋 {disp_type} Document")
    with col_badge:
        st.success("✓ Ready")

    with st.container():
        st.markdown('<div class="output-card">', unsafe_allow_html=True)
        st.markdown(content)
        st.markdown('</div>', unsafe_allow_html=True)

    st.markdown("---")
    st.markdown("#### 📥 Download")
    col_dl1, col_dl2 = st.columns(2)

    doc = Document()
    doc.add_heading(f"{disp_type} — {label_disp}", 0)
    doc.add_paragraph(f"Generated by AI Document Generator · {disp_type}")
    doc.add_paragraph("")
    for line in content.split('\n'):
        if line.strip():
            if   line.startswith('# '):   doc.add_heading(line.replace('# ',  '').strip(), 1)
            elif line.startswith('## '):  doc.add_heading(line.replace('## ', '').strip(), 2)
            elif line.startswith('### '): doc.add_heading(line.replace('### ','').strip(), 3)
            else:                          doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    with col_dl1:
        st.download_button(
            label="📄 Download as Word (.docx)",
            data=buffer,
            file_name=f"{disp_type}_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
    with col_dl2:
        st.download_button(
            label="📝 Download as Text (.txt)",
            data=content,
            file_name=f"{disp_type}_document.txt",
            mime="text/plain",
            use_container_width=True
        )